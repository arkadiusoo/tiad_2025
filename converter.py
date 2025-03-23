import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as convert_pdf
import os
import platform

def convert_xlsx_to_docx(xlsx_path,
                         docx_path,
                         fileType,
                         headers,
                         headers_bold,
                         font_size,
                         font,
                         form,
                         original_name,
                         alignment,
                         spacing):
    font_size = int(font_size)
    spacing = Pt(spacing)

    df = pd.read_excel(xlsx_path)
    doc = Document()

    # Nagłówek dokumentu
    heading = doc.add_heading(original_name, level=1)
    for run in heading.runs:
        run.font.name = font
        run.font.size = Pt(font_size)

    columns_per_table = 10

    if form == 1:
        if df.shape[1] > 10:
            for start_col in range(0, df.shape[1], columns_per_table):
                end_col = min(start_col + columns_per_table, df.shape[1])
                sub_df = df.iloc[:, start_col:end_col]

                table = doc.add_table(rows=sub_df.shape[0] + 1, cols=sub_df.shape[1])
                table.style = 'Table Grid'

                if headers:
                    for j, column in enumerate(sub_df.columns):
                        cell = table.cell(0, j)
                        para = cell.paragraphs[0]
                        para.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                        para.paragraph_format.space_after = spacing
                        run = para.add_run(str(column))
                        run.font.name = font
                        run.font.size = Pt(font_size)
                        if headers_bold:
                            run.bold = True

                for i, row in sub_df.iterrows():
                    for j, value in enumerate(row):
                        cell = table.cell(i + 1, j)
                        para = cell.paragraphs[0]
                        para.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                        para.paragraph_format.space_after = spacing
                        run = para.add_run(str(value))
                        run.font.name = font
                        run.font.size = Pt(font_size)

                doc.add_paragraph("")
        else:
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'

            if headers:
                for j, column in enumerate(df.columns):
                    cell = table.cell(0, j)
                    para = cell.paragraphs[0]
                    para.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                    para.paragraph_format.space_after = spacing
                    run = para.add_run(str(column))
                    run.font.name = font
                    run.font.size = Pt(font_size)
                    if headers_bold:
                        run.bold = True

            for i, row in df.iterrows():
                for j, value in enumerate(row):
                    cell = table.cell(i + 1, j)
                    para = cell.paragraphs[0]
                    para.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                    para.paragraph_format.space_after = spacing
                    run = para.add_run(str(value))
                    run.font.name = font
                    run.font.size = Pt(font_size)

    elif form == 2:
        if headers:
            header_line = " ".join([str(col) for col in df.columns])
            p = doc.add_paragraph()
            p.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
            p.paragraph_format.space_after = spacing
            run = p.add_run(header_line.strip())
            run.font.name = font
            run.font.size = Pt(font_size)
            if headers_bold:
                run.bold = True

        for _, row in df.iterrows():
            line_text = " ".join([str(value) for value in row])
            p = doc.add_paragraph()
            p.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
            p.paragraph_format.space_after = spacing
            run = p.add_run(line_text.strip())
            run.font.name = font
            run.font.size = Pt(font_size)

    elif form == 3 or form == 4:
        for idx, row in df.iterrows():
            # Numeracja
            p_num = doc.add_paragraph()
            p_num.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
            p_num.paragraph_format.space_after = spacing
            run_num = p_num.add_run(f"{idx + 1}.")
            run_num.font.name = font
            run_num.font.size = Pt(font_size)
            if headers_bold:
                run_num.bold = True

            for col in df.columns:
                value = str(row[col])
                p = doc.add_paragraph()
                p.alignment = {"Lewo": WD_ALIGN_PARAGRAPH.LEFT, "Środek": WD_ALIGN_PARAGRAPH.CENTER, "Prawo": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                p.paragraph_format.space_after = spacing

                if headers:
                    run1 = p.add_run(f"{col}: ")
                    run1.font.name = font
                    run1.font.size = Pt(font_size)
                    if headers_bold:
                        run1.bold = True

                    run2 = p.add_run(value)
                    run2.font.name = font
                    run2.font.size = Pt(font_size)
                else:
                    run = p.add_run(value)
                    run.font.name = font
                    run.font.size = Pt(font_size)

            if form == 4 and idx < len(df) - 1:
                doc.add_page_break()

    # Zapis dokumentu
    doc.save(docx_path)

    # Eksport PDF (jeśli wybrano)
    if fileType:
        try:
            convert_pdf(docx_path)
            os.remove(docx_path)
        except Exception as e:
            print("Błąd konwersji do PDF:", e)
            if platform.system() == "Linux":
                print("Konwersja PDF nie jest obsługiwana na Linuxie przez docx2pdf.")
